"""解析《排名汇总.docx》→ data/major_ranks/<year>.json，供 import_major_ranks 入库。

文档结构：每所学校一段标题 + 一张表（专业/年份/等级/全国排名）。
覆盖 2023–2026，专业：软件工程、人工智能、计算机科学与技术。

用法（在 backend 目录）:
  python scripts/parse_rank_summary_docx.py
  python scripts/parse_rank_summary_docx.py --docx "D:/UGit/data/排名汇总.docx" --import
"""

from __future__ import annotations

import argparse
import json
import sys
from collections import defaultdict
from pathlib import Path
from typing import Any

BACKEND = Path(__file__).resolve().parents[1]
SCRIPTS = Path(__file__).resolve().parent
sys.path.insert(0, str(BACKEND))
sys.path.insert(0, str(SCRIPTS))

OUT_DIR = BACKEND / "data" / "major_ranks"
DEFAULT_DOCX = Path(r"D:\UGit\data\排名汇总.docx")

SELF_SCHOOL = "广东财经大学"

# 文档内院校 → 对标池划分（与排名汇总口径一致）
REGIONAL_POOL = {
    "广东财经大学",
    "广东外语外贸大学",
    "广州大学",
    "广东技术师范大学",
}
FINANCE_POOL = {
    "广东财经大学",
    "南京审计大学",
    "天津财经大学",
    "山西财经大学",
    "吉林财经大学",
    "重庆工商大学",
}


def _s(v: Any) -> str:
    return str(v or "").strip()


def _as_int(v: Any) -> int | None:
    t = _s(v).replace(",", "").replace("，", "")
    if not t or t in {"-", "—", "－", "/", "无"}:
        return None
    try:
        return int(float(t))
    except ValueError:
        return None


def _grade(v: Any) -> str | None:
    t = _s(v).upper().replace(" ", "")
    if not t or t in {"-", "—", "－", "/"}:
        return None
    return t


def parse_docx(path: Path) -> dict[str, list[dict[str, Any]]]:
    """返回 school → [{major, year, grade, nationalRank}, ...]。"""
    try:
        from docx import Document
    except ImportError as exc:  # pragma: no cover
        raise SystemExit("需要 python-docx：pip install python-docx") from exc

    doc = Document(str(path))
    # 按文档顺序：段落学校名与紧随其后的表格一一对应
    # python-docx 不保证 body 顺序 API 简单，用段落索引 + 表格序号启发式：
    # 实际文档：学校名段落 →（空段）→ 下一张表
    schools: list[str] = []
    for p in doc.paragraphs:
        t = _s(p.text)
        if t and ("大学" in t or (t.endswith("学院") and len(t) < 20)):
            schools.append(t)

    if len(schools) != len(doc.tables):
        # 回退：仅保留与表数量相同的前 N 个校名
        print(
            f"[warn] schools({len(schools)}) != tables({len(doc.tables)}); "
            f"align by min length"
        )
        n = min(len(schools), len(doc.tables))
        schools = schools[:n]
        tables = doc.tables[:n]
    else:
        tables = doc.tables

    by_school: dict[str, list[dict[str, Any]]] = {}
    for school, table in zip(schools, tables):
        rows: list[dict[str, Any]] = []
        last_major = ""
        for ri, row in enumerate(table.rows):
            cells = [_s(c.text).replace("\n", "") for c in row.cells]
            if ri == 0:
                continue  # header
            while len(cells) < 4:
                cells.append("")
            major, year_s, grade_s, rank_s = cells[0], cells[1], cells[2], cells[3]
            if major:
                last_major = major
            else:
                major = last_major
            year = _as_int(year_s)
            if not major or year is None:
                continue
            rows.append(
                {
                    "major": major,
                    "year": year,
                    "grade": _grade(grade_s),
                    "nationalRank": _as_int(rank_s),
                }
            )
        by_school[school] = rows
        print(f"[parse] {school}: {len(rows)} rows")
    return by_school


def _pool_ranks(
    by_school: dict[str, list[dict[str, Any]]],
    *,
    major: str,
    year: int,
    pool: set[str],
) -> list[dict[str, Any]]:
    """某专业某年在对标池内的全国名次列表（有排名才入选）。"""
    out: list[dict[str, Any]] = []
    for school, rows in by_school.items():
        if school not in pool:
            continue
        hit = next(
            (
                r
                for r in rows
                if r["major"] == major and r["year"] == year and r["nationalRank"] is not None
            ),
            None,
        )
        if not hit:
            continue
        item: dict[str, Any] = {"school": school, "rank": int(hit["nationalRank"])}
        if school == SELF_SCHOOL:
            item["isSelf"] = True
        out.append(item)
    out.sort(key=lambda x: x["rank"])
    return out


def _self_rank_in_pool(peers: list[dict[str, Any]]) -> int | None:
    for i, p in enumerate(peers, start=1):
        if p.get("isSelf"):
            return i
    return None


def build_year_payloads(
    by_school: dict[str, list[dict[str, Any]]],
) -> dict[int, dict[str, Any]]:
    self_rows = by_school.get(SELF_SCHOOL) or []
    if not self_rows:
        raise SystemExit(f"未找到本校「{SELF_SCHOOL}」表格")

    # major → year → row
    self_index: dict[str, dict[int, dict[str, Any]]] = defaultdict(dict)
    for r in self_rows:
        self_index[r["major"]][r["year"]] = r

    years = sorted({r["year"] for r in self_rows})
    payloads: dict[int, dict[str, Any]] = {}

    for year in years:
        majors_out: list[dict[str, Any]] = []
        for major, by_year in sorted(self_index.items()):
            cur = by_year.get(year)
            if not cur or cur.get("nationalRank") is None:
                # 本校该年无排名则跳过该专业
                continue
            prev = by_year.get(year - 1)
            yoy = None
            if prev and prev.get("nationalRank") is not None:
                # 名次变小 = 上升，正数表示上升位数
                yoy = int(prev["nationalRank"]) - int(cur["nationalRank"])

            peer_schools = _pool_ranks(
                by_school, major=major, year=year, pool=REGIONAL_POOL
            )
            finance_peers = _pool_ranks(
                by_school, major=major, year=year, pool=FINANCE_POOL
            )
            majors_out.append(
                {
                    "name": major,
                    "grade": cur.get("grade"),
                    "nationalRank": cur["nationalRank"],
                    "yoyChange": yoy,
                    "provincialRank": _self_rank_in_pool(peer_schools),
                    "financePeerRank": _self_rank_in_pool(finance_peers),
                    "softDimensions": [],
                    "peerSchools": peer_schools,
                    "financePeerSchools": finance_peers,
                }
            )

        payloads[year] = {
            "year": year,
            "source": "rank-summary-docx",
            "selfSchool": SELF_SCHOOL,
            "majors": majors_out,
            "meta": {
                "regionalPool": sorted(REGIONAL_POOL),
                "financePool": sorted(FINANCE_POOL),
                "docxSchools": list(by_school.keys()),
            },
        }
    return payloads


def write_payloads(payloads: dict[int, dict[str, Any]], out_dir: Path) -> list[Path]:
    out_dir.mkdir(parents=True, exist_ok=True)
    written: list[Path] = []
    for year, payload in sorted(payloads.items()):
        path = out_dir / f"{year}.json"
        path.write_text(
            json.dumps(payload, ensure_ascii=False, indent=2),
            encoding="utf-8",
        )
        written.append(path)
        print(
            f"[write] {path.name} majors={len(payload['majors'])} "
            f"names={[m['name'] for m in payload['majors']]}"
        )
    return written


def update_peer_yaml() -> None:
    """把文档对标池写回 peer_schools.yaml，便于后续 bcmr 切片对齐。"""
    yaml_path = SCRIPTS / "config" / "peer_schools.yaml"
    text = (
        "# 对标院校名单（与《排名汇总.docx》口径对齐；采集器按全国榜切片时使用）\n"
        f"selfSchool: {SELF_SCHOOL}\n"
        "\n"
        "# 综合 / 省内对标池（含本校）——文档内有排名的广东高校\n"
        "regional:\n"
        + "".join(f"  - {s}\n" for s in sorted(REGIONAL_POOL, key=lambda x: (x != SELF_SCHOOL, x)))
        + "\n"
        "# 财经院校对标池（含本校）——文档内财经类院校\n"
        "finance:\n"
        + "".join(f"  - {s}\n" for s in sorted(FINANCE_POOL, key=lambda x: (x != SELF_SCHOOL, x)))
    )
    # 把本校放到列表靠后更直观：上面 sort 已让 self 靠后
    yaml_path.write_text(text, encoding="utf-8")
    print(f"[write] {yaml_path}")


def main() -> None:
    parser = argparse.ArgumentParser(description="解析排名汇总.docx → major_ranks JSON")
    parser.add_argument("--docx", type=Path, default=DEFAULT_DOCX)
    parser.add_argument("--out-dir", type=Path, default=OUT_DIR)
    parser.add_argument(
        "--import",
        dest="do_import",
        action="store_true",
        help="生成后执行 import_ugit_data --only major_ranks",
    )
    parser.add_argument(
        "--no-yaml",
        action="store_true",
        help="不更新 peer_schools.yaml",
    )
    args = parser.parse_args()

    if not args.docx.is_file():
        raise SystemExit(f"找不到文件: {args.docx}")

    by_school = parse_docx(args.docx)
    payloads = build_year_payloads(by_school)
    write_payloads(payloads, args.out_dir)
    if not args.no_yaml:
        update_peer_yaml()

    if args.do_import:
        import subprocess

        cmd = [
            sys.executable,
            str(SCRIPTS / "import_ugit_data.py"),
            "--data-root",
            str(BACKEND / "data"),
            "--only",
            "major_ranks",
        ]
        print("[import]", " ".join(cmd))
        subprocess.check_call(cmd, cwd=str(BACKEND))


if __name__ == "__main__":
    main()
